"""
주간 미션 파일을 로드하고 파싱하는 모듈.

.ipynb 문제 파일과 채점 기준표(.xlsx, .docx)를 로드합니다.
⚠️ 정답 파일은 제외합니다.
"""

import re
import sys
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

# 프로젝트 루트 및 notebooks 폴더 경로 추가
PROJECT_ROOT = Path(__file__).parent.parent.parent
NOTEBOOKS_DIR = Path(__file__).parent.parent / "notebooks"
sys.path.insert(0, str(NOTEBOOKS_DIR))

# notebooks 모듈의 클래스 직접 import
from notebook_loader import NotebookLoader, ParsedNotebook


class MissionType(str, Enum):
    """미션 파일 타입."""

    PROBLEM = "문제"  # 문제 노트북
    SOLUTION = "정답"  # 정답 노트북 (제외 대상)
    RUBRIC = "채점기준"  # 채점 기준표


@dataclass
class RubricItem:
    """채점 기준 항목."""

    criteria: str  # 평가 기준
    points: str  # 배점
    description: str  # 설명


@dataclass
class ParsedMission:
    """
    파싱된 미션을 나타내는 클래스.

    Attributes:
        file_path: 파일 경로
        mission_type: 미션 타입 (문제/정답/채점기준)
        course: 과목명
        week: 주차 (예: "w1", "w4")
        mission_name: 미션명
        instructor: 마스터명
        notebook: 파싱된 노트북 (ipynb인 경우)
        rubric_items: 채점 기준 항목 (채점기준표인 경우)
        raw_text: 원본 텍스트 (docx/xlsx인 경우)
    """

    file_path: Path
    mission_type: MissionType
    course: str = ""
    week: str = ""
    mission_name: str = ""
    instructor: str = ""
    notebook: ParsedNotebook | None = None
    rubric_items: list[RubricItem] = field(default_factory=list)
    raw_text: str = ""

    @property
    def is_problem(self) -> bool:
        """문제 파일인지 확인."""
        return self.mission_type == MissionType.PROBLEM

    @property
    def is_rubric(self) -> bool:
        """채점 기준표인지 확인."""
        return self.mission_type == MissionType.RUBRIC


class MissionLoader:
    """
    주간 미션 파일을 로드하는 클래스.

    문제 노트북과 채점 기준표를 로드합니다.
    ⚠️ 정답 파일은 자동으로 제외됩니다.

    예시:
        ```python
        loader = MissionLoader()
        missions = loader.load_from_directory("weekly_mission/")

        for mission in missions:
            if mission.is_problem:
                print(f"📝 문제: {mission.mission_name}")
            elif mission.is_rubric:
                print(f"📋 채점기준: {mission.mission_name}")
        ```
    """

    # 정답 파일 패턴 (제외 대상)
    SOLUTION_PATTERNS = [
        r"\(정답\)",
        r"정답\.ipynb$",
        r"_정답",
        r"\(해설\)",
    ]

    # 문제 파일 패턴
    PROBLEM_PATTERNS = [
        r"\(문제\)",
        r"문제\.ipynb$",
        r"_문제",
    ]

    # 채점 기준표 패턴
    RUBRIC_PATTERNS = [
        r"채점.*기준",
        r"rubric",
        r"grading",
    ]

    # 주차 패턴
    WEEK_PATTERN = re.compile(r"w(\d+)", re.IGNORECASE)

    # 폴더명 패턴 (마스터명 추출)
    FOLDER_PATTERN = re.compile(r"^\d+\.\s*(?P<course>.+?)\s*\((?P<instructor>\w+)\s*마스터\)$")

    def __init__(self, verbose: bool = False) -> None:
        """
        MissionLoader 초기화.

        Args:
            verbose: 상세 로그 출력 여부
        """
        self.verbose = verbose
        self.notebook_loader = NotebookLoader()

    def load_from_directory(
        self,
        directory: Path | str,
        include_solutions: bool = False,
    ) -> list[ParsedMission]:
        """
        디렉토리 내 미션 파일을 로드합니다.

        Args:
            directory: 미션 파일이 있는 디렉토리
            include_solutions: 정답 파일 포함 여부 (기본: False)

        Returns:
            ParsedMission 리스트
        """
        directory = Path(directory)

        if not directory.exists():
            raise FileNotFoundError(f"디렉토리를 찾을 수 없습니다: {directory}")

        missions: list[ParsedMission] = []

        # 노트북 파일 (.ipynb)
        for ipynb_file in directory.rglob("*.ipynb"):
            mission = self._load_notebook(ipynb_file)
            if mission:
                # 정답 파일 필터링
                if mission.mission_type == MissionType.SOLUTION and not include_solutions:
                    if self.verbose:
                        print(f"   ⏭️ 정답 제외: {ipynb_file.name}")
                    continue
                missions.append(mission)

        # 채점 기준표 (.xlsx)
        for xlsx_file in directory.rglob("*.xlsx"):
            mission = self._load_xlsx(xlsx_file)
            if mission:
                missions.append(mission)

        # 채점 기준표 (.docx)
        for docx_file in directory.rglob("*.docx"):
            mission = self._load_docx(docx_file)
            if mission:
                missions.append(mission)

        if self.verbose:
            print(f"📂 로드 완료: {len(missions)}개 미션 파일")

        return missions

    def _load_notebook(self, file_path: Path) -> ParsedMission | None:
        """노트북 파일을 로드합니다."""
        try:
            notebook = self.notebook_loader.load_from_file(file_path)

            # 미션 타입 결정
            mission_type = self._detect_mission_type(file_path.name)

            # 메타데이터 추출
            course, instructor = self._extract_course_instructor(file_path)
            week = self._extract_week(file_path.name)
            mission_name = self._extract_mission_name(file_path.name)

            if self.verbose:
                type_icon = "📝" if mission_type == MissionType.PROBLEM else "📖"
                print(f"   {type_icon} {file_path.name}")

            return ParsedMission(
                file_path=file_path,
                mission_type=mission_type,
                course=course,
                week=week,
                mission_name=mission_name,
                instructor=instructor,
                notebook=notebook,
            )

        except Exception as e:
            if self.verbose:
                print(f"   ⚠️ 로드 실패: {file_path.name} - {e}")
            return None

    def _load_xlsx(self, file_path: Path) -> ParsedMission | None:
        """Excel 채점 기준표를 로드합니다."""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheet = wb.active

            # 모든 셀 텍스트 추출
            text_parts: list[str] = []
            for row in sheet.iter_rows():
                row_texts = []
                for cell in row:
                    if cell.value:
                        row_texts.append(str(cell.value))
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

            raw_text = "\n".join(text_parts)

            # 메타데이터 추출
            course, instructor = self._extract_course_instructor(file_path)
            mission_name = self._extract_mission_name(file_path.name)

            if self.verbose:
                print(f"   📋 {file_path.name}")

            return ParsedMission(
                file_path=file_path,
                mission_type=MissionType.RUBRIC,
                course=course,
                mission_name=mission_name,
                instructor=instructor,
                raw_text=raw_text,
            )

        except ImportError:
            if self.verbose:
                print("   ⚠️ openpyxl 필요: pip install openpyxl")
            return None
        except Exception as e:
            if self.verbose:
                print(f"   ⚠️ 로드 실패: {file_path.name} - {e}")
            return None

    def _load_docx(self, file_path: Path) -> ParsedMission | None:
        """Word 채점 기준표를 로드합니다."""
        try:
            import docx

            doc = docx.Document(file_path)

            # 모든 단락 텍스트 추출
            text_parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 테이블 텍스트도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))

            raw_text = "\n".join(text_parts)

            # 메타데이터 추출
            course, instructor = self._extract_course_instructor(file_path)
            mission_name = self._extract_mission_name(file_path.name)

            if self.verbose:
                print(f"   📋 {file_path.name}")

            return ParsedMission(
                file_path=file_path,
                mission_type=MissionType.RUBRIC,
                course=course,
                mission_name=mission_name,
                instructor=instructor,
                raw_text=raw_text,
            )

        except ImportError:
            if self.verbose:
                print("   ⚠️ python-docx 필요: pip install python-docx")
            return None
        except Exception as e:
            if self.verbose:
                print(f"   ⚠️ 로드 실패: {file_path.name} - {e}")
            return None

    def _detect_mission_type(self, filename: str) -> MissionType:
        """파일명에서 미션 타입을 감지합니다."""
        # 정답 파일 체크
        for pattern in self.SOLUTION_PATTERNS:
            if re.search(pattern, filename, re.IGNORECASE):
                return MissionType.SOLUTION

        # 채점 기준표 체크
        for pattern in self.RUBRIC_PATTERNS:
            if re.search(pattern, filename, re.IGNORECASE):
                return MissionType.RUBRIC

        # 기본: 문제 파일
        return MissionType.PROBLEM

    def _extract_course_instructor(self, file_path: Path) -> tuple[str, str]:
        """폴더명에서 과목과 마스터명을 추출합니다."""
        for parent in file_path.parents:
            match = self.FOLDER_PATTERN.match(parent.name)
            if match:
                course = match.group("course").strip()
                instructor = match.group("instructor").strip()
                return course, instructor
        return "", ""

    def _extract_week(self, filename: str) -> str:
        """파일명에서 주차를 추출합니다."""
        match = self.WEEK_PATTERN.search(filename)
        if match:
            return f"w{match.group(1)}"
        return ""

    def _extract_mission_name(self, filename: str) -> str:
        """파일명에서 미션명을 추출합니다."""
        # 확장자 제거
        name = Path(filename).stem

        # 패턴 제거
        patterns_to_remove = [
            r"\[.*?\]",  # [Pytorch] 등
            r"\(문제\)|\(정답\)",  # (문제), (정답)
            r"w\d+_",  # w1_ 등
            r"채점.*기준.*",  # 채점 기준표
            r"_$",  # 끝 언더스코어
        ]

        for pattern in patterns_to_remove:
            name = re.sub(pattern, "", name, flags=re.IGNORECASE)

        return name.strip().strip("_").strip()


# =============================================================================
# CLI 테스트
# =============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("사용법: python mission_loader.py <directory>")
        sys.exit(1)

    directory = sys.argv[1]
    loader = MissionLoader(verbose=True)

    try:
        missions = loader.load_from_directory(directory)

        print(f"\n{'=' * 60}")
        print(f"📂 총 {len(missions)}개 미션 파일 로드")
        print("=" * 60)

        problems = [m for m in missions if m.is_problem]
        rubrics = [m for m in missions if m.is_rubric]

        print(f"📝 문제: {len(problems)}개")
        print(f"📋 채점기준: {len(rubrics)}개")

    except Exception as e:
        print(f"❌ 오류: {e}")
        sys.exit(1)
